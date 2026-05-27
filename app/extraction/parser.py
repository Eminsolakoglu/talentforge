from pathlib import Path
from typing import Dict, Any
import logging
import io

import fitz  # PyMuPDF
import pytesseract
from PIL import Image

logger = logging.getLogger(__name__)


class CVParser:
    """CV dosyalarını parse eden temel sınıf"""

    def __init__(self):
        logger.info("✅ CVParser initialized (pdfplumber + python-docx + OCR)")

    def parse(self, file_path: str | Path) -> Dict[str, Any]:
        """
        CV dosyasını parse eder.
        PDF için pdfplumber -> pypdf -> pdfminer -> OCR fallback zinciri kullanılır.
        DOCX için python-docx kullanılır.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"CV dosyası bulunamadı: {file_path}")

        logger.info(f"📄 Parsing CV: {file_path.name}")

        try:
            suffix = file_path.suffix.lower()

            if suffix == ".pdf":
                return self._parse_pdf(file_path)

            if suffix == ".docx":
                return self._parse_docx(file_path)

            raise ValueError(f"Desteklenmeyen dosya türü: {file_path.suffix}")

        except Exception as e:
            logger.error(f"Parsing error: {e}")
            raise

    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """PDF dosyalarını parse eder"""

        # 1. pdfplumber
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                pages_text = [page.extract_text() or "" for page in pdf.pages]

            full_text = "\n\n".join(pages_text).strip()

            if len(full_text) > 100:
                return self._pdf_result(
                    file_path=file_path,
                    text=full_text,
                    pages=len(pages_text),
                    parser="pdfplumber"
                )

            logger.warning("⚠️ pdfplumber boş metin döndürdü, pypdf deneniyor...")

        except Exception as e:
            logger.warning(f"⚠️ pdfplumber hatası: {e}")

        # 2. pypdf
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            pages_text = [page.extract_text() or "" for page in reader.pages]
            full_text = "\n\n".join(pages_text).strip()

            if len(full_text) > 100:
                return self._pdf_result(
                    file_path=file_path,
                    text=full_text,
                    pages=len(pages_text),
                    parser="pypdf"
                )

            logger.warning("⚠️ pypdf de boş döndürdü, pdfminer deneniyor...")

        except Exception as e:
            logger.warning(f"⚠️ pypdf hatası: {e}")

        # 3. pdfminer
        try:
            from pdfminer.high_level import extract_text

            full_text = extract_text(str(file_path)).strip()

            if len(full_text) > 100:
                return self._pdf_result(
                    file_path=file_path,
                    text=full_text,
                    pages=0,
                    parser="pdfminer"
                )

            logger.warning("⚠️ pdfminer de boş döndürdü, OCR deneniyor...")

        except Exception as e:
            logger.warning(f"⚠️ pdfminer hatası: {e}")

        # 4. OCR image based PDF
        try:
            doc = fitz.open(str(file_path))
            all_text = []

            for page_index, page in enumerate(doc, start=1):
                pix = page.get_pixmap(dpi=250)
                img = Image.open(io.BytesIO(pix.tobytes("png")))

                text = pytesseract.image_to_string(
                    img,
                    lang="eng+tur",
                    config="--psm 6"
                )

                all_text.append(text)

            doc.close()

            full_text = "\n\n".join(all_text).strip()

            if len(full_text) > 100:
                return self._pdf_result(
                    file_path=file_path,
                    text=full_text,
                    pages=len(all_text),
                    parser="ocr-tesseract"
                )

            logger.warning("⚠️ OCR da boş döndürdü")

        except Exception as e:
            logger.warning(f"⚠️ OCR hatası: {e}")

        return self._pdf_result(
            file_path=file_path,
            text="",
            pages=0,
            parser="none"
        )

    def _pdf_result(
        self,
        file_path: Path,
        text: str,
        pages: int,
        parser: str
    ) -> Dict[str, Any]:
        """PDF parse sonucunu standart formatta döndürür"""

        logger.info(f"✅ PDF parsed with {parser}: {len(text)} karakter")

        return {
            "file_name": file_path.name,
            "file_type": "pdf",
            "raw_text": text[:15000],
            "total_pages": pages,
            "status": "success" if text else "empty",
            "parser": parser,
        }

    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        from docx import Document
        doc = Document(file_path)
    
        parts = []
    
        # Normal paragraflar
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
    
        # Tablolar (çoğu CV tabloyla yapılmış)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
    
        full_text = "\n".join(parts)
    
        return {
            "file_name": file_path.name,
            "file_type": "docx",
            "raw_text": full_text[:15000],
            "total_paragraphs": len(parts),
            "status": "success" if full_text else "empty",
            "parser": "python-docx",
        }